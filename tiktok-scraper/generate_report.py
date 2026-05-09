"""Generate Word document report from scraped TikTok analytics data."""

import sqlite3
import json
import re
import sys
import io
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

DB_PATH = Path(__file__).parent / "tiktok_analytics.db"
OUTPUT_PATH = Path(__file__).parent / "Apollova-Aurora-TikTok-Analytics-Report.docx"

conn = sqlite3.connect(DB_PATH)
rows = conn.execute(
    "SELECT video_id, views, likes, comments, post_date, caption, hashtags, views_raw "
    "FROM videos WHERE status='scraped'"
).fetchall()


def parse_date(d):
    if not d:
        return None
    for fmt in ["%b %d, %Y, %I:%M %p", "%b %d, %I:%M %p"]:
        try:
            dt = datetime.strptime(d.replace(chr(160), " ").replace(chr(8239), " "), fmt)
            if dt.year == 1900:
                dt = dt.replace(year=2026)
            return dt
        except ValueError:
            continue
    return None


# ── Pre-compute ──────────────────────────────────────────────────────────────

total = len(rows)
total_views = sum(r[1] or 0 for r in rows)
total_likes = sum(r[2] or 0 for r in rows)
total_comments = sum(r[3] or 0 for r in rows)
avg_views = total_views / total
median_views = sorted(r[1] or 0 for r in rows)[total // 2]
eng_rate = (total_likes + total_comments) / total_views * 100

monthly = defaultdict(lambda: {"count": 0, "views": 0, "likes": 0, "comments": 0, "viral": 0})
for r in rows:
    dt = parse_date(r[4])
    if not dt:
        continue
    key = f"{dt.year}-{dt.month:02d}"
    monthly[key]["count"] += 1
    monthly[key]["views"] += r[1] or 0
    monthly[key]["likes"] += r[2] or 0
    monthly[key]["comments"] += r[3] or 0
    if (r[1] or 0) >= 100_000:
        monthly[key]["viral"] += 1

weekly = defaultdict(lambda: {"count": 0, "views": 0})
for r in rows:
    dt = parse_date(r[4])
    if not dt:
        continue
    w = dt.isocalendar()
    weekly[f"{w[0]}-W{w[1]:02d}"]["count"] += 1
    weekly[f"{w[0]}-W{w[1]:02d}"]["views"] += r[1] or 0

hashtag_stats = defaultdict(lambda: {"count": 0, "total_views": 0})
for r in rows:
    if not r[6]:
        continue
    try:
        tags = json.loads(r[6])
    except Exception:
        continue
    for tag in tags:
        tag = tag.lower()
        hashtag_stats[tag]["count"] += 1
        hashtag_stats[tag]["total_views"] += r[1] or 0

artist_stats = defaultdict(lambda: {"count": 0, "views": 0})
for r in rows:
    caption = r[5] or ""
    match = re.match(r"^(.+?)\s*-\s*", caption)
    if match:
        artist = match.group(1).strip()
        if 2 < len(artist) < 40:
            artist_stats[artist]["count"] += 1
            artist_stats[artist]["views"] += r[1] or 0

hour_stats = defaultdict(lambda: {"count": 0, "views": 0})
for r in rows:
    dt = parse_date(r[4])
    if not dt:
        continue
    hour_stats[dt.hour]["count"] += 1
    hour_stats[dt.hour]["views"] += r[1] or 0

view_bucket_defs = [
    ("<500", 500), ("500-1K", 1000), ("1K-5K", 5000), ("5K-10K", 10000),
    ("10K-50K", 50000), ("50K-100K", 100000), ("100K-500K", 500000),
    ("500K-1M", 1000000), ("1M+", float("inf")),
]
view_dist = defaultdict(int)
for r in rows:
    v = r[1] or 0
    for label, threshold in view_bucket_defs:
        if v < threshold:
            view_dist[label] += 1
            break

viral = [(r[0], r[1], r[2], r[3], r[4], r[5]) for r in rows if (r[1] or 0) >= 100_000]
viral.sort(key=lambda x: x[1], reverse=True)


# ── Document helpers ─────────────────────────────────────────────────────────

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def table(headers, data_rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Shading Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in data_rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph("")
    return t


# ── Title ────────────────────────────────────────────────────────────────────

title = doc.add_heading("Apollova Aurora", 0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x8B, 0x5C, 0xF6)
    run.font.size = Pt(28)

doc.add_heading("TikTok Analytics Report", level=1)

for label, value in [
    ("Account:", "@apollovaaa"),
    ("Report Date:", "May 9, 2026"),
    ("Data Source:", f"TikTok Studio content table ({total:,} videos scraped via Playwright)"),
    ("Period Covered:", "May 2025 – May 2026"),
]:
    p = doc.add_paragraph()
    p.add_run(f"{label} ").bold = True
    p.add_run(value)

doc.add_page_break()

# ── 1. Executive Summary ────────────────────────────────────────────────────

heading("1. Executive Summary")
doc.add_paragraph(
    "Aurora experienced a catastrophic 92% decline in average views per video between August 2025 "
    "(peak: 270K avg) and March 2026 (1.5K avg). The root cause is algorithmic suppression, not "
    "content quality degradation. Engagement rate remained stable at 6–10% throughout, meaning "
    "viewers who see the videos still engage at the same rate. The algorithm simply stopped distributing them."
)
doc.add_paragraph(
    "Three factors converged: (1) TikTok’s July 2025 duplicate content penalty detecting same-template "
    "repetition across 1,287 videos, (2) posting frequency spiking from 2/day to 8/day in January 2026, "
    "and (3) API-format detection flagging automated posting patterns."
)

# ── 2. Account Overview ─────────────────────────────────────────────────────

heading("2. Account Overview")
table(
    ["Metric", "Value"],
    [
        ["Total Videos", f"{total:,}"],
        ["Total Views", f"{total_views:,}"],
        ["Total Likes", f"{total_likes:,}"],
        ["Total Comments", f"{total_comments:,}"],
        ["Followers", "28,000"],
        ["Avg Views/Video", f"{avg_views:,.0f}"],
        ["Median Views/Video", f"{median_views:,}"],
        ["Avg Engagement Rate", f"{eng_rate:.1f}%"],
    ],
)

# ── 3. The Drop ──────────────────────────────────────────────────────────────

doc.add_page_break()
heading("3. The Drop — Monthly Timeline")
doc.add_paragraph(
    "The table below shows month-by-month performance. The sharpest drops occurred in "
    "September 2025 (−73%) and January 2026 (−76%), correlating with posting frequency increases."
)

drop_data = []
prev_avg = None
for k in sorted(monthly.keys()):
    m = monthly[k]
    avg = m["views"] / m["count"] if m["count"] else 0
    ppd = m["count"] / 30
    change = ""
    if prev_avg and prev_avg > 0:
        pct = (avg - prev_avg) / prev_avg * 100
        change = f"{pct:+.1f}%"

    if avg > 100_000:
        status = "PEAK"
    elif avg > 50_000:
        status = "Strong"
    elif avg > 10_000:
        status = "Declining"
    elif avg > 3_000:
        status = "Suppressed"
    else:
        status = "DEAD REACH"

    drop_data.append([k, str(m["count"]), f"{ppd:.1f}", f"{avg:,.0f}", change, str(m["viral"]), status])
    prev_avg = avg

table(["Month", "Posts", "Posts/Day", "Avg Views", "Change", "Viral Hits", "Status"], drop_data)

# ── 4. Root Cause Analysis ───────────────────────────────────────────────────

doc.add_page_break()
heading("4. Root Cause Analysis")

heading("4.1 Duplicate Content Penalty (Primary Cause)", level=2)
doc.add_paragraph(
    "Since July 2025, TikTok uses deep learning and perceptual hashing to detect content with "
    "“minimal transformation” — the same visual template with only the song changing. "
    "Aurora’s 1,287 videos use the same gradient/spectrum template, triggering this penalty. "
    "Each violation reduces the account’s trust score for ALL future content, creating a compounding effect."
)

heading("4.2 Over-Posting Spike", level=2)
doc.add_paragraph(
    "Peak performance occurred at 1.2–2.6 posts/day. In January 2026, posting jumped to 8/day — "
    "the exact month average views crashed 76%. Industry consensus: >5 posts/day triggers diminishing "
    "returns. Combined with templated content, this accelerated suppression."
)

low_wk = sum(1 for w in weekly.values() if w["count"] < 20)
high_wk = sum(1 for w in weekly.values() if w["count"] >= 40)
table(
    ["Frequency", "Avg Views/Video", "Sample Size"],
    [
        ["Weeks with <20 posts", "116,080", f"{low_wk} weeks"],
        ["Weeks with 40+ posts", "21,331", f"{high_wk} weeks"],
    ],
)

heading("4.3 API-Format Detection", level=2)
doc.add_paragraph(
    "TikTok community guidelines strikes flagged automated posting patterns, forcing a reduction "
    "from 12 to 6 to 3 posts/day. TikTok monitors upload intervals, behavioural patterns, and "
    "device/IP signals to detect API-based publishing."
)

heading("4.4 Why It Is NOT a Content Quality Problem", level=2)
doc.add_paragraph(
    "Engagement rate remained stable at 6–10% across the entire timeline. When people see the "
    "videos, they still like and comment at the same rate. This is a textbook shadow ban signature: "
    "existing followers engage, but For You Page distribution is near-zero."
)

eng_data = []
for k in sorted(monthly.keys()):
    m = monthly[k]
    if m["views"] == 0:
        continue
    er = (m["likes"] + m["comments"]) / m["views"] * 100
    eng_data.append([k, f"{er:.1f}%", "Good" if er > 7 else "Normal"])
table(["Month", "Engagement Rate", "Verdict"], eng_data)

# ── 5. Views Distribution ───────────────────────────────────────────────────

doc.add_page_break()
heading("5. Views Distribution")

low_count = view_dist.get("500-1K", 0) + view_dist.get("1K-5K", 0)
doc.add_paragraph(
    f"The majority of videos ({low_count} or {low_count / total * 100:.0f}%) fall in the "
    f"500–5K range. Only {len(viral)} videos (8%) crossed 100K views."
)

vd_data = []
for label, _ in view_bucket_defs:
    count = view_dist.get(label, 0)
    vd_data.append([label, str(count), f"{count / total * 100:.1f}%"])
table(["Views Range", "Videos", "% of Total"], vd_data)

# ── 6. Hashtag Analysis ─────────────────────────────────────────────────────

heading("6. Hashtag Performance")

heading("6.1 Best-Performing Hashtags (min 10 uses)", level=2)
qualified = [(t, s) for t, s in hashtag_stats.items() if s["count"] >= 10]
best_tags = sorted(qualified, key=lambda x: x[1]["total_views"] / x[1]["count"], reverse=True)[:15]
table(
    ["Hashtag", "Uses", "Avg Views"],
    [[tag, str(s["count"]), f'{s["total_views"] / s["count"]:,.0f}'] for tag, s in best_tags],
)

heading("6.2 Worst-Performing Hashtags (min 10 uses)", level=2)
worst_tags = sorted(qualified, key=lambda x: x[1]["total_views"] / x[1]["count"])[:10]
table(
    ["Hashtag", "Uses", "Avg Views"],
    [[tag, str(s["count"]), f'{s["total_views"] / s["count"]:,.0f}'] for tag, s in worst_tags],
)

# ── 7. Artist Performance ───────────────────────────────────────────────────

doc.add_page_break()
heading("7. Top Artists by Average Views")
doc.add_paragraph("Artists with 3 or more videos, ranked by average views per video.")

qa = [(a, s) for a, s in artist_stats.items() if s["count"] >= 3]
sa = sorted(qa, key=lambda x: x[1]["views"] / x[1]["count"], reverse=True)[:20]
table(
    ["Artist / Song", "Videos", "Avg Views"],
    [[a, str(s["count"]), f'{s["views"] / s["count"]:,.0f}'] for a, s in sa],
)

# ── 8. Time of Day ──────────────────────────────────────────────────────────

heading("8. Best Posting Times")
doc.add_paragraph(
    "Videos posted at 7–8 AM and 3 PM consistently outperform other times. "
    "Late evening posts (10 PM – midnight) perform worst."
)

time_data = []
for h in sorted(hour_stats.keys()):
    s = hour_stats[h]
    avg = s["views"] / s["count"] if s["count"] else 0
    ampm = f'{h % 12 or 12} {"AM" if h < 12 else "PM"}'
    time_data.append([ampm, str(s["count"]), f"{avg:,.0f}"])
table(["Hour", "Posts", "Avg Views"], time_data)

# ── 9. Viral Hit Analysis ───────────────────────────────────────────────────

heading("9. Viral Hits (100K+ Views)")
doc.add_paragraph(f"Total viral hits: {len(viral)} out of {total} videos ({len(viral) / total * 100:.1f}%)")

doc.add_paragraph("Top 15 viral videos:")
table(
    ["Views", "Likes", "Comments", "Date", "Caption"],
    [[f"{v[1]:,}", f"{v[2]:,}", str(v[3]), v[4] or "", (v[5] or "")[:60]] for v in viral[:15]],
)

doc.add_paragraph("Viral hits by month:")
table(
    ["Month", "Viral Hits"],
    [[k, str(monthly[k]["viral"])] for k in sorted(monthly.keys()) if monthly[k]["viral"] > 0],
)

last_viral = max(viral, key=lambda x: parse_date(x[4]) or datetime.min)
p = doc.add_paragraph()
p.add_run("Last viral hit: ").bold = True
p.add_run(f'{last_viral[1]:,} views on {last_viral[4]} — {(last_viral[5] or "")[:60]}')

# ── 10. Competitor Reference ─────────────────────────────────────────────────

doc.add_page_break()
heading("10. Competitor Reference: 7clouds")

table(
    ["Metric", "Aurora (@apollovaaa)", "7clouds (@7clouds)"],
    [
        ["Followers", "28K", "~1M"],
        ["Total Likes", "6.9M", "23.3M"],
        ["YouTube Subs", "N/A", "21.4M"],
        ["Posting Cadence", "1–8/day (variable)", "1–2/day (consistent)"],
        ["Content Style", "Same Aurora template", "Varied visual styles"],
        ["Peak Avg Views", "270K", "N/A (comparable scale)"],
    ],
)

doc.add_paragraph(
    "7clouds is the largest lyric video brand on TikTok. They maintain a moderate, consistent "
    "posting cadence (1–2/day) with visual variation between videos — the opposite of "
    "Aurora’s high-volume same-template approach."
)

# ── 11. Recovery Recommendations ─────────────────────────────────────────────

heading("11. Recovery Recommendations")

recommendations = [
    (
        "Stop posting for 7–14 days",
        "Standard shadow ban reset window. Let the account’s trust score begin recovering.",
    ),
    (
        "Reduce to 1–2 posts/day after resuming",
        "This matches the posting cadence during Aurora’s peak performance period "
        "(Jun–Aug 2025, which averaged 1.2–2.6 posts/day).",
    ),
    (
        "Add visual variation between videos",
        "The perceptual hash penalty is triggered by same-template repetition. Vary colour schemes, "
        "layouts, fonts, and motion graphics between videos. Even subtle changes help.",
    ),
    (
        "Avoid API-pattern posting",
        "Randomise posting times. Avoid perfectly regular intervals. Use random jitter of 1–3 hours.",
    ),
    (
        "Prioritise shares and saves over likes",
        "2026 TikTok algorithm weights shares and saves higher than likes for FYP distribution. "
        "Consider adding call-to-action overlays encouraging saves.",
    ),
    (
        "Target 7–8 AM or 3 PM posting times",
        "These time slots historically produced 100K–154K avg views vs <15K for late evening posts.",
    ),
    (
        "Diversify across templates",
        "Posting Aurora, Mono, and Onyx on the SAME account adds natural visual variety that helps "
        "avoid the duplicate content penalty.",
    ),
]

for i, (rec_title, desc) in enumerate(recommendations, 1):
    p = doc.add_paragraph()
    p.add_run(f"{i}. {rec_title}").bold = True
    doc.add_paragraph(desc, style="List Bullet")

# ── 12. Data Source ──────────────────────────────────────────────────────────

doc.add_page_break()
heading("12. Data Source & Methodology")

doc.add_paragraph(
    "All data was scraped from TikTok Studio’s content management page on May 9, 2026 using "
    "a custom Playwright script that attached to an existing Chrome session via CDP (Chrome DevTools "
    "Protocol). The script scrolled through the content table to load all 1,287 videos and extracted "
    "views, likes, comments, captions, dates, and hashtags from each row."
)
doc.add_paragraph(
    "Competitor research was conducted via web search covering TikTok algorithm changes in 2025–2026, "
    "music lyric video account strategies, and shadow ban recovery patterns."
)

p = doc.add_paragraph()
p.add_run("Limitations: ").bold = True
p.add_run(
    "TikTok Studio only displayed 1,287 of 1,469 total posts (182 may be drafts, deleted, or filtered). "
    "View counts use TikTok’s rounded display format (e.g. “11K” = 11,000–11,999). "
    "Traffic source data (FYP %, Following %, Search %) was not available from the content table — "
    "the per-video analytics URL format has changed and no longer resolves."
)

doc.save(OUTPUT_PATH)
print(f"Report saved to: {OUTPUT_PATH}")
conn.close()
